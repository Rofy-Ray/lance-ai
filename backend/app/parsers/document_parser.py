import os
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
from docx import Document
from pathlib import Path
from typing import Dict, Any, List
import asyncio
import logging

logger = logging.getLogger(__name__)

class DocumentParser:
    """Document parsing with OCR support for various file types"""
    
    def __init__(self):
        # Configure Tesseract path if needed
        tesseract_cmd = os.getenv("TESSERACT_CMD")
        if tesseract_cmd:
            pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    
    async def parse_document(self, file_path: str) -> Dict[str, Any]:
        """Parse document and extract text content with page information"""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"Document not found: {file_path}")
            
            # Determine file type and parse accordingly
            if file_path.suffix.lower() == '.pdf':
                return await self._parse_pdf(file_path)
            elif file_path.suffix.lower() in ['.doc', '.docx']:
                return await self._parse_docx(file_path)
            elif file_path.suffix.lower() == '.txt':
                return await self._parse_text(file_path)
            elif file_path.suffix.lower() in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                return await self._parse_image_ocr(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_path.suffix}")
                
        except Exception as e:
            logger.error(f"Failed to parse document {file_path}: {e}")
            return {
                "text": "",
                "pages": [],
                "error": str(e),
                "file_type": file_path.suffix.lower() if hasattr(file_path, 'suffix') else "unknown"
            }
    
    async def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF with fallback to OCR for scanned documents"""
        try:
            doc = fitz.open(file_path)
            pages = []
            full_text = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                
                # If no text extracted, try OCR
                if not text.strip():
                    try:
                        pix = page.get_pixmap()
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        text = pytesseract.image_to_string(img)
                    except Exception as ocr_error:
                        logger.warning(f"OCR failed on page {page_num + 1}: {ocr_error}")
                        text = ""
                
                page_info = {
                    "page_number": page_num + 1,
                    "text": text,
                    "word_count": len(text.split()),
                    "char_count": len(text)
                }
                pages.append(page_info)
                full_text.append(text)
            
            doc.close()
            
            return {
                "text": "\n\n".join(full_text),
                "pages": pages,
                "page_count": len(pages),
                "file_type": "pdf",
                "total_words": sum(page["word_count"] for page in pages),
                "parsing_method": "text_extraction_with_ocr_fallback"
            }
            
        except Exception as e:
            raise Exception(f"PDF parsing failed: {str(e)}")
    
    async def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX document"""
        try:
            doc = Document(file_path)
            text_blocks = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_blocks.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_blocks.append(" | ".join(row_text))
            
            full_text = "\n".join(text_blocks)
            
            # Estimate pages (roughly 250 words per page)
            words = full_text.split()
            words_per_page = 250
            estimated_pages = max(1, len(words) // words_per_page + (1 if len(words) % words_per_page else 0))
            
            pages = []
            for page_num in range(estimated_pages):
                start_word = page_num * words_per_page
                end_word = min((page_num + 1) * words_per_page, len(words))
                page_text = " ".join(words[start_word:end_word])
                
                pages.append({
                    "page_number": page_num + 1,
                    "text": page_text,
                    "word_count": len(page_text.split()),
                    "char_count": len(page_text)
                })
            
            return {
                "text": full_text,
                "pages": pages,
                "page_count": len(pages),
                "file_type": "docx",
                "total_words": len(words),
                "parsing_method": "docx_extraction"
            }
            
        except Exception as e:
            raise Exception(f"DOCX parsing failed: {str(e)}")
    
    async def _parse_text(self, file_path: Path) -> Dict[str, Any]:
        """Parse plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            
            # Split into estimated pages
            words = text.split()
            words_per_page = 250
            estimated_pages = max(1, len(words) // words_per_page + (1 if len(words) % words_per_page else 0))
            
            pages = []
            for page_num in range(estimated_pages):
                start_word = page_num * words_per_page
                end_word = min((page_num + 1) * words_per_page, len(words))
                page_text = " ".join(words[start_word:end_word])
                
                pages.append({
                    "page_number": page_num + 1,
                    "text": page_text,
                    "word_count": len(page_text.split()),
                    "char_count": len(page_text)
                })
            
            return {
                "text": text,
                "pages": pages,
                "page_count": len(pages),
                "file_type": "text",
                "total_words": len(words),
                "parsing_method": "text_file"
            }
            
        except Exception as e:
            raise Exception(f"Text parsing failed: {str(e)}")
    
    async def _parse_image_ocr(self, file_path: Path) -> Dict[str, Any]:
        """Parse image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            if not text.strip():
                text = "[Image contains no readable text]"
            
            words = text.split()
            
            return {
                "text": text,
                "pages": [{
                    "page_number": 1,
                    "text": text,
                    "word_count": len(words),
                    "char_count": len(text)
                }],
                "page_count": 1,
                "file_type": "image",
                "total_words": len(words),
                "parsing_method": "ocr"
            }
            
        except Exception as e:
            raise Exception(f"Image OCR failed: {str(e)}")
    
    def extract_quotes_with_location(self, parsed_doc: Dict[str, Any], quote_text: str, context_chars: int = 50) -> List[Dict[str, Any]]:
        """Find quote locations within parsed document"""
        matches = []
        
        for page in parsed_doc.get("pages", []):
            page_text = page["text"]
            page_number = page["page_number"]
            
            # Find all occurrences of the quote in this page
            start = 0
            while True:
                pos = page_text.lower().find(quote_text.lower(), start)
                if pos == -1:
                    break
                
                # Calculate line number (rough estimate)
                lines_before = page_text[:pos].count('\n')
                line_number = lines_before + 1
                
                # Extract context
                context_start = max(0, pos - context_chars)
                context_end = min(len(page_text), pos + len(quote_text) + context_chars)
                context = page_text[context_start:context_end]
                
                matches.append({
                    "page": page_number,
                    "line_range": f"{line_number}-{line_number}",
                    "quote_span": quote_text,
                    "context": context,
                    "position": pos
                })
                
                start = pos + 1
        
        return matches
